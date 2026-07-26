import asyncio
import json
import re
import shutil
import threading
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from app.core.config import settings


_WHITESPACE_PATTERN = re.compile(r"\s+")
_ocr_instance: Any | None = None
_ocr_lock = threading.Lock()
_ocr_predict_lock = threading.Lock()


@dataclass(slots=True, frozen=True)
class ExtractedText:
    text: str
    extraction_method: str
    page_count: int | None = None


class TextExtractor:
    """根据扩展名提取 PDF 或 DOCX 文本，阻塞工作统一放在线程池执行。"""

    SUPPORTED_SUFFIXES = {".pdf", ".docx"}

    async def extract(self, file_path: Path) -> ExtractedText:
        suffix = file_path.suffix.lower()
        if suffix == ".pdf":
            return await asyncio.to_thread(self._extract_pdf, file_path)
        if suffix == ".docx":
            return await asyncio.to_thread(self._extract_docx, file_path)
        raise ValueError("暂不支持该文件格式，目前仅支持 PDF 和 DOCX")

    @staticmethod
    def is_supported(filename: str) -> bool:
        return Path(filename).suffix.lower() in TextExtractor.SUPPORTED_SUFFIXES

    def _extract_pdf(self, file_path: Path) -> ExtractedText:
        import fitz

        try:
            with fitz.open(file_path) as document:
                page_count = document.page_count
                page_texts = [page.get_text("text") for page in document]
        except Exception as exc:
            raise ValueError("PDF 文件无法打开或内容已损坏") from exc

        text = self._normalize_text("\n\n".join(page_texts))
        effective_chars = len(_WHITESPACE_PATTERN.sub("", text))
        if effective_chars >= settings.pdf_text_min_chars:
            return ExtractedText(
                text=text,
                extraction_method="pymupdf",
                page_count=page_count,
            )

        # 有效文字过少时按扫描件处理，逐页渲染后交给复用的 PaddleOCR 实例。
        try:
            ocr_text = self._extract_scanned_pdf(file_path)
        except Exception as exc:
            raise ValueError(
                "扫描版 PDF OCR 识别失败，请检查 PaddleOCR 模型和运行环境"
            ) from exc
        if not ocr_text:
            raise ValueError("扫描版 PDF 未识别出有效文字")
        return ExtractedText(
            text=ocr_text,
            extraction_method="paddleocr",
            page_count=page_count,
        )

    def _extract_scanned_pdf(self, file_path: Path) -> str:
        import fitz

        ocr = _get_ocr_instance()
        page_texts: list[str] = []
        render_dir = file_path.parent / f"{file_path.stem}_ocr_pages"
        render_dir.mkdir(parents=True, exist_ok=True)

        try:
            with fitz.open(file_path) as document:
                matrix = fitz.Matrix(
                    settings.ocr_render_scale,
                    settings.ocr_render_scale,
                )
                for page_number, page in enumerate(document):
                    image_path = render_dir / f"page_{page_number + 1:04d}.png"
                    pixmap = page.get_pixmap(matrix=matrix, alpha=False)
                    pixmap.save(image_path)
                    page_texts.append(_run_ocr(ocr, image_path))
        finally:
            shutil.rmtree(render_dir, ignore_errors=True)

        return self._normalize_text("\n\n".join(page_texts))

    def _extract_docx(self, file_path: Path) -> ExtractedText:
        from docx import Document

        try:
            document = Document(file_path)
        except Exception as exc:
            raise ValueError("DOCX 文件无法打开或内容已损坏") from exc

        fragments = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                cells = [self._normalize_text(cell.text) for cell in row.cells]
                fragments.append(" | ".join(cell for cell in cells if cell))

        text = self._normalize_text("\n".join(fragments))
        if not text:
            raise ValueError("DOCX 文件中没有可提取的文字")
        return ExtractedText(text=text, extraction_method="python-docx")

    @staticmethod
    def _normalize_text(text: str) -> str:
        lines = [
            _WHITESPACE_PATTERN.sub(" ", line).strip()
            for line in text.replace("\x00", "").splitlines()
        ]
        return "\n".join(line for line in lines if line)


def _get_ocr_instance() -> Any:
    """延迟初始化 OCR，避免启动应用时加载大模型，并在进程内全局复用。"""
    global _ocr_instance
    if _ocr_instance is not None:
        return _ocr_instance

    with _ocr_lock:
        if _ocr_instance is not None:
            return _ocr_instance

        from paddleocr import PaddleOCR

        try:
            # PaddleOCR 3.x 参数，关闭与本阶段无关的文档预处理模型。
            _ocr_instance = PaddleOCR(
                lang=settings.ocr_language,
                use_doc_orientation_classify=False,
                use_doc_unwarping=False,
                use_textline_orientation=False,
                enable_mkldnn=False,
            )
        except (TypeError, ValueError):
            # 兼容仍使用 2.x API 的部署环境。
            _ocr_instance = PaddleOCR(
                lang=settings.ocr_language,
                use_angle_cls=True,
                show_log=False,
                enable_mkldnn=False,
            )
        return _ocr_instance


def _run_ocr(ocr: Any, image_path: Path) -> str:
    """兼容 PaddleOCR 2.x 的 ocr() 与 3.x 的 predict() 返回结构。"""
    # Paddle 推理器不是线程安全对象，多个请求共享单例时需要串行执行。
    with _ocr_predict_lock:
        if hasattr(ocr, "predict"):
            raw_result = ocr.predict(input=str(image_path))
            texts = _extract_v3_texts(raw_result)
        else:
            raw_result = ocr.ocr(str(image_path), cls=True)
            texts = _extract_v2_texts(raw_result)
    return "\n".join(text for text in texts if text.strip())


def _extract_v3_texts(raw_result: Any) -> list[str]:
    texts: list[str] = []
    for item in raw_result or []:
        payload: Any = item
        json_value = getattr(item, "json", None)
        if json_value is not None:
            payload = json_value() if callable(json_value) else json_value
        if isinstance(payload, str):
            try:
                payload = json.loads(payload)
            except json.JSONDecodeError:
                continue
        texts.extend(_find_rec_texts(payload))
    return texts


def _find_rec_texts(payload: Any) -> list[str]:
    if isinstance(payload, dict):
        direct = payload.get("rec_texts")
        if isinstance(direct, list):
            return [str(value) for value in direct]
        texts: list[str] = []
        for value in payload.values():
            texts.extend(_find_rec_texts(value))
        return texts
    if isinstance(payload, list):
        texts = []
        for value in payload:
            texts.extend(_find_rec_texts(value))
        return texts
    return []


def _extract_v2_texts(raw_result: Any) -> list[str]:
    texts: list[str] = []
    for page_result in raw_result or []:
        for line in page_result or []:
            if not isinstance(line, (list, tuple)) or len(line) < 2:
                continue
            recognition = line[1]
            if (
                isinstance(recognition, (list, tuple))
                and recognition
                and isinstance(recognition[0], str)
            ):
                texts.append(recognition[0])
    return texts
